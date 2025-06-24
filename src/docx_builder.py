"""
DOCX Builder Module
Creates Microsoft Word documents from layout data while preserving formatting.
"""

import logging
from pathlib import Path
from typing import Dict, List, Any, Optional, Tuple
import io

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from PIL import Image


class DocxBuilder:
    """Builds DOCX documents from layout data."""
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
        self.document = None
        self.current_section = None
    
    def create_document(self, layout_data, output_path: Path) -> bool:
        """
        Create a DOCX document from layout data.
        
        Args:
            layout_data: DocumentLayout object
            output_path: Path for the output DOCX file
            
        Returns:
            True if successful, False otherwise
        """
        try:
            self.logger.info("Creating DOCX document...")
            
            # Initialize document
            self.document = Document()
            
            # Set up document properties
            self._setup_document_properties(layout_data)
            
            # Process each page
            for page_idx, page_layout in enumerate(layout_data.pages):
                self.logger.debug(f"Processing page {page_idx + 1}")
                self._process_page(page_layout, layout_data)
                
                # Add page break between pages (except for the last page)
                if page_idx < len(layout_data.pages) - 1:
                    self._add_page_break()
            
            # Save document
            self.document.save(str(output_path))
            self.logger.info(f"DOCX document saved: {output_path}")
            return True
            
        except Exception as e:
            self.logger.error(f"Error creating DOCX document: {str(e)}")
            return False
    
    def _setup_document_properties(self, layout_data) -> None:
        """Set up document-wide properties."""
        # Get the first page to set initial page size
        if layout_data.pages:
            first_page = layout_data.pages[0]
            
            # Set page size and margins
            section = self.document.sections[0]
            
            # Convert PDF points to inches (72 points = 1 inch)
            page_width_inches = first_page.page_size[0] / 72
            page_height_inches = first_page.page_size[1] / 72
            
            section.page_width = Inches(page_width_inches)
            section.page_height = Inches(page_height_inches)
            
            # Set margins
            margins = first_page.margins
            section.top_margin = Inches(margins.get('top', 72) / 72)
            section.bottom_margin = Inches(margins.get('bottom', 72) / 72)
            section.left_margin = Inches(margins.get('left', 72) / 72)
            section.right_margin = Inches(margins.get('right', 72) / 72)
            
            self.current_section = section
    
    def _process_page(self, page_layout, layout_data) -> None:
        """Process a single page and add its content to the document."""
        
        # Add headers first
        if page_layout.headers:
            self._add_headers(page_layout.headers, layout_data)
        
        # Process main content elements in reading order
        reading_order_elements = self._get_reading_order(page_layout.elements)
        
        for element in reading_order_elements:
            self._add_element_to_document(element, layout_data)
        
        # Add footers last
        if page_layout.footers:
            self._add_footers(page_layout.footers, layout_data)
    
    def _get_reading_order(self, elements: List) -> List:
        """Sort elements in reading order (top to bottom, left to right)."""
        return sorted(elements, key=lambda e: (e.bbox[1], e.bbox[0]))
    
    def _add_element_to_document(self, element, layout_data) -> None:
        """Add a layout element to the document."""
        
        if element.element_type == 'text':
            self._add_text_element(element, layout_data)
        elif element.element_type == 'image':
            self._add_image_element(element)
        elif element.element_type == 'ocr_text':
            self._add_ocr_text_element(element, layout_data)
    
    def _add_text_element(self, element, layout_data) -> None:
        """Add a text element to the document."""
        text_block = element.content
        
        # Create paragraph
        paragraph = self.document.add_paragraph()
        run = paragraph.add_run(text_block.text)
        
        # Apply formatting
        self._apply_text_formatting(run, text_block, layout_data)
        
        # Set paragraph alignment based on position
        self._set_paragraph_alignment(paragraph, element.bbox)
    
    def _add_image_element(self, element) -> None:
        """Add an image element to the document."""
        try:
            image_block = element.content
            
            # Convert image data to PIL Image
            if image_block.image_format == "png":
                image = Image.open(io.BytesIO(image_block.image_data))
            else:
                # Handle raw image data - convert to bytes and create image
                image_bytes = io.BytesIO(image_block.image_data)
                image = Image.open(image_bytes)
            
            # Save image to temporary bytes buffer
            img_buffer = io.BytesIO()
            image.save(img_buffer, format='PNG')
            img_buffer.seek(0)
            
            # Calculate display size (convert from PDF points to inches)
            width_inches = (image_block.bbox[2] - image_block.bbox[0]) / 72
            height_inches = (image_block.bbox[3] - image_block.bbox[1]) / 72
            
            # Add image to document
            paragraph = self.document.add_paragraph()
            run = paragraph.add_run()
            run.add_picture(img_buffer, width=Inches(width_inches), height=Inches(height_inches))
            
            # Center the image
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        except Exception as e:
            self.logger.warning(f"Error adding image: {str(e)}")
            # Add placeholder text if image fails
            paragraph = self.document.add_paragraph()
            run = paragraph.add_run("[Image could not be inserted]")
            run.italic = True
    
    def _add_ocr_text_element(self, element, layout_data) -> None:
        """Add OCR extracted text to the document."""
        ocr_result = element.content
        
        # Create paragraph with OCR text
        paragraph = self.document.add_paragraph()
        run = paragraph.add_run(ocr_result.text)
        
        # Mark as OCR text with special formatting
        run.italic = True
        run.font.color.rgb = RGBColor(128, 128, 128)  # Gray color
        
        # Add comment or note about OCR confidence
        if hasattr(ocr_result, 'confidence') and ocr_result.confidence < 70:
            comment_run = paragraph.add_run(f" [OCR confidence: {ocr_result.confidence:.0f}%]")
            comment_run.font.size = Pt(8)
            comment_run.font.color.rgb = RGBColor(255, 0, 0)  # Red for low confidence
    
    def _apply_text_formatting(self, run, text_block, layout_data) -> None:
        """Apply formatting to a text run based on the original text block."""
        
        # Map font name
        font_name = self._map_font_name(text_block.font_name, layout_data.font_mapping)
        run.font.name = font_name
        
        # Set font size
        run.font.size = Pt(text_block.font_size)
        
        # Apply font styles based on flags
        if text_block.font_flags & 2**4:  # Bold flag
            run.bold = True
        if text_block.font_flags & 2**6:  # Italic flag
            run.italic = True
        
        # Set text color
        if hasattr(text_block, 'color') and text_block.color != 0:
            color_rgb = self._convert_color(text_block.color)
            if color_rgb:
                run.font.color.rgb = RGBColor(*color_rgb)
    
    def _map_font_name(self, pdf_font_name: str, font_mapping: Dict[str, str]) -> str:
        """Map PDF font name to Word-compatible font."""
        # Try exact match first
        if pdf_font_name in font_mapping:
            return font_mapping[pdf_font_name]
        
        # Try partial matches
        for pdf_font, word_font in font_mapping.items():
            if pdf_font.lower() in pdf_font_name.lower():
                return word_font
        
        # Default fallback
        if 'Arial' in pdf_font_name or 'Helvetica' in pdf_font_name:
            return 'Arial'
        elif 'Times' in pdf_font_name:
            return 'Times New Roman'
        elif 'Courier' in pdf_font_name:
            return 'Courier New'
        else:
            return 'Arial'  # Ultimate fallback
    
    def _convert_color(self, color_int: int) -> Optional[Tuple[int, int, int]]:
        """Convert PDF color integer to RGB tuple."""
        try:
            # PDF color is typically in format 0xRRGGBB
            r = (color_int >> 16) & 0xFF
            g = (color_int >> 8) & 0xFF
            b = color_int & 0xFF
            return (r, g, b)
        except:
            return None
    
    def _set_paragraph_alignment(self, paragraph, bbox: Tuple[float, float, float, float]) -> None:
        """Set paragraph alignment based on position."""
        # This is a simplified alignment detection
        # In a real implementation, you would analyze the position relative to page width
        # For now, just use left alignment as default
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    def _add_headers(self, header_elements: List, layout_data) -> None:
        """Add header elements to the document."""
        if not header_elements:
            return
        
        # Create header section
        header = self.current_section.header
        header_paragraph = header.paragraphs[0]
        header_paragraph.clear()
        
        for element in header_elements:
            if element.element_type == 'text':
                run = header_paragraph.add_run(element.content.text + " ")
                self._apply_text_formatting(run, element.content, layout_data)
    
    def _add_footers(self, footer_elements: List, layout_data) -> None:
        """Add footer elements to the document."""
        if not footer_elements:
            return
        
        # Create footer section
        footer = self.current_section.footer
        footer_paragraph = footer.paragraphs[0]
        footer_paragraph.clear()
        
        for element in footer_elements:
            if element.element_type == 'text':
                run = footer_paragraph.add_run(element.content.text + " ")
                self._apply_text_formatting(run, element.content, layout_data)
    
    def _add_page_break(self) -> None:
        """Add a page break to the document."""
        self.document.add_page_break()
    
    def _add_table(self, table_data: Dict[str, Any]) -> None:
        """Add a table to the document (placeholder for future enhancement)."""
        # This would be implemented to handle table structures
        # detected in the PDF layout
        pass
    
    def _preserve_spacing(self, element, previous_element) -> None:
        """Preserve spacing between elements (placeholder for future enhancement)."""
        # This would calculate and apply appropriate spacing
        # between document elements based on their positions
        pass
    
    def set_document_metadata(self, title: str = None, author: str = None, subject: str = None) -> None:
        """Set document metadata."""
        if self.document:
            core_props = self.document.core_properties
            if title:
                core_props.title = title
            if author:
                core_props.author = author
            if subject:
                core_props.subject = subject
    
    def optimize_document(self) -> None:
        """Optimize the document structure (placeholder for future enhancement)."""
        # This could include:
        # - Merging consecutive paragraphs with same formatting
        # - Optimizing image sizes
        # - Removing empty paragraphs
        # - Adjusting spacing for better layout
        pass
